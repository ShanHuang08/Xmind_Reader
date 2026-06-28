"""Parse Confluence-exported Word documents into structured text."""

from __future__ import annotations

import email
import re
from email import policy
from pathlib import Path
from typing import Any

from lxml import etree, html


class DocReadError(Exception):
    """Raised when a vendor source document cannot be read."""


def parse_vendor_doc(path: Path) -> dict[str, Any]:
    path = Path(path)
    if not path.exists():
        raise DocReadError(f"Document does not exist: {path}")
    if path.suffix.lower() not in {".doc", ".docx", ".html", ".htm"}:
        raise DocReadError(f"Unsupported document extension: {path.suffix}")

    raw_bytes = path.read_bytes()
    if path.suffix.lower() == ".docx":
        return _parse_docx(path)

    html_text = _extract_html_text(raw_bytes)
    if not html_text:
        raise DocReadError(
            "Could not find HTML content. If this is a binary .doc file, export it as .docx or HTML."
        )
    return _parse_html_document(path, html_text)


def _parse_docx(path: Path) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocReadError("python-docx is required to read .docx files.") from exc

    document = Document(path)
    paragraphs = [
        {"style": paragraph.style.name if paragraph.style else "", "text": paragraph.text.strip()}
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
    text = "\n".join(item["text"] for item in paragraphs)
    return {
        "source_file": path.name,
        "source_path": str(path),
        "format": "docx",
        "title": _title_from_paragraphs(paragraphs) or path.stem,
        "paragraphs": paragraphs,
        "tables": tables,
        "links": [],
        "plain_text": text,
    }


def _extract_html_text(raw_bytes: bytes) -> str:
    sample = raw_bytes[:512].decode("utf-8", errors="ignore").lower()
    if "content-type: multipart/" in sample or "mime-version:" in sample:
        message = email.message_from_bytes(raw_bytes, policy=policy.default)
        for part in message.walk():
            content_type = part.get_content_type()
            if content_type == "text/html":
                payload = part.get_payload(decode=True) or b""
                charset = part.get_content_charset() or "utf-8"
                return _decode_payload(payload, charset)
    text = raw_bytes.decode("utf-8", errors="ignore")
    if "<html" in text.lower() or "<body" in text.lower():
        return text
    return ""


def _decode_payload(payload: bytes, charset: str) -> str:
    normalized = _normalize_charset(charset)
    candidates = [normalized, "utf-8", "utf-16", "cp950"]
    seen = set()
    for candidate in candidates:
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)
        try:
            return payload.decode(candidate, errors="replace")
        except LookupError:
            continue
    return payload.decode("utf-8", errors="replace")


def _normalize_charset(charset: str) -> str:
    normalized = (charset or "utf-8").strip().strip('"').lower()
    aliases = {
        "unicode": "utf-16",
        "utf8": "utf-8",
        "utf16": "utf-16",
    }
    return aliases.get(normalized, normalized)


def _parse_html_document(path: Path, html_text: str) -> dict[str, Any]:
    root = html.fromstring(html_text)
    title = _clean(" ".join(root.xpath("//title/text()"))) or path.stem
    headings = []
    paragraphs = []
    for element in root.xpath("//h1|//h2|//h3|//h4|//h5|//h6|//p|//li"):
        text = _clean(element.text_content())
        if not text:
            continue
        tag = element.tag.lower()
        if tag.startswith("h"):
            headings.append({"level": int(tag[1]), "text": text})
            paragraphs.append({"style": tag, "text": text})
        else:
            paragraphs.append({"style": tag, "text": text})

    tables = []
    tables_detailed = []
    for table in root.xpath("//table"):
        rows = []
        detailed_rows = []
        for row in table.xpath(".//tr"):
            cells = []
            detailed_cells = []
            for cell in row.xpath("./th|./td"):
                text = _clean(cell.text_content())
                checkbox = _checkbox_state(cell)
                cells.append(text if checkbox is None else checkbox)
                detailed_cells.append(
                    {
                        "text": text,
                        "checkbox": checkbox,
                        "tag": cell.tag.lower(),
                    }
                )
            if any(cells):
                rows.append(cells)
                detailed_rows.append(detailed_cells)
        if rows:
            tables.append(rows)
            tables_detailed.append(detailed_rows)

    links = [
        {"text": _clean(link.text_content()), "href": link.get("href", "")}
        for link in root.xpath("//a[@href]")
    ]
    plain_text = "\n".join(item["text"] for item in paragraphs)
    return {
        "source_file": path.name,
        "source_path": str(path),
        "format": "confluence-html-doc",
        "title": title,
        "headings": headings,
        "paragraphs": paragraphs,
        "tables": tables,
        "tables_detailed": tables_detailed,
        "links": links,
        "plain_text": plain_text,
    }


def _title_from_paragraphs(paragraphs: list[dict[str, str]]) -> str:
    for paragraph in paragraphs:
        if paragraph.get("style", "").lower().startswith("heading"):
            return paragraph.get("text", "")
    return paragraphs[0]["text"] if paragraphs else ""


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _checkbox_state(cell: etree._Element) -> str | None:
    task_items = cell.xpath(
        ".//*[contains(concat(' ', normalize-space(@class), ' '), ' inline-task-list ')]//li"
    )
    task_items.extend(
        cell.xpath(".//li[contains(concat(' ', normalize-space(@class), ' '), ' checked ')]")
    )
    if not task_items:
        return None
    return "checked" if any(_has_class(item, "checked") for item in task_items) else "unchecked"


def _has_class(element: etree._Element, class_name: str) -> bool:
    classes = (element.get("class") or "").split()
    return class_name in classes
